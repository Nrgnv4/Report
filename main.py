import matplotlib
import matplotlib.pyplot as plt
from Comtrade import ComtradeRecord
import os
import json
import logging
import re
import random
from docx import Document
from docx.shared import Cm

# import settings as s
matplotlib.use('Agg')


logging.basicConfig(
    format=u'%(levelname)-8s [%(filename)-12s:%(lineno)-4d # %(asctime)s] %(message)s',
    level=logging.INFO, filename='sample.log')

logging.info('Start!')


choice = {
    'Ie': 'SetIe',
    'If': 'SetIf',
    'Fg': '',
    'Ug': 'SetUg',
    'Pg': '',
    'Qg': '',
    'Alpha': 'Fg',

}


class Csv():
    """docstring for Csv"""

    def __init__(self, path):
        self.file = []
        self.signals = {}
        with open(path, 'r') as f:
            for i in f:
                self.file.append(i.replace('\n', '').replace(',', '.'))
        self.split_data()
        self.get_signals()

    def split_data(self):
        data = [self.line_split(i) for i in self.file]
        self.data = data[1:]

    def get_signals(self):
        key = 0
        for signal in self.data[0][1:]:
            key += 1
            self.signals.update({key: signal.strip()})

    def line_split(self, line):
        a = line.split(';')
        return a

    def get_time(self):
        time = [float(i[0]) for i in self.data[1:]]
        return time

    def get_channel(self, num):
        data = [float(i[num]) for i in self.data[1:]]
        return data


class Scope():
    """docstring for Scope"""

    def __init__(self, choice, path_to_file, convert_to_csv = False):
        self.choice = choice

        # чтение настроек
        self.global_settings = Helper().read_settings()

        # получение времени, анaлоговых и дискретных сигналов из файла
        self.get_data_from_file(path_to_file)

        if convert_to_csv:
            # преобразование в Csv
            self.convert_to_csv(path_to_file)
        else:
            # преобразование словаря имен и номеров аналоговых сигналов
            self.transform_analogs_names()

            # формирование выбранных сигналов self.choosen_signals
            self.analogs_select(choice)

            # формирование макета осциллограммы
            self.prepare_figure()

            # формирование и размещение на макете всех сигналов
            self.add_analogs()
            if self.comtrade:
                self.add_digitals()

            # добавление подписи по "y"
            plt.xlabel('t,сек', fontsize=7, rotation=0)

            # plt.savefig(f'{path_to_save}{s.IMAGE_FILENAME}',
            #     transparent=self.global_settings['IMAGE_TRANSPARENT'])

            # сохранение файла в png
            plt.savefig(path_to_file[:-3],
                        transparent=self.global_settings['IMAGE_TRANSPARENT'])
            plt.close()


    def get_data_from_file(self, path_to_file):
        """В зависимости от типа файла получает сигналы из файла
        """
        self.comtrade = False
        self.csv = False
        if path_to_file[-3:] == 'cfg':
            comtrade = ComtradeRecord(path_to_file)
            comtrade.ReadDataFile()
            self.comtrade = comtrade
            self.analogs = {i:j.rstrip() for i,j in comtrade.signals['Analog'].items()}
            self.digitals = comtrade.signals['Digital']
            self.time = comtrade.getTime()
            if self.global_settings['DELETE_EMPTY_DIGITAL_SIGNALS']:
                self.del_empty_digital_signals()

        if path_to_file[-3:] == 'csv':
            csv = Csv(path_to_file)
            self.csv = csv
            self.analogs = csv.signals
            self.time = csv.get_time()
            self.digitals = []

    def analog_data(self, channel):
        if self.comtrade:
            return self.comtrade.getAnalogChannelData(channel)
        if self.csv:
            return self.csv.get_channel(channel)

    def convert_to_csv(self, path_to_file):
        import numpy as np
        import pandas as pd 
        first_row = ['time']
        first_row.extend(self.analogs.values())
        first_row = np.array(first_row)
        data = []
        [data.append(self.analog_data(key)) for key in self.analogs.keys()]
        data = np.hstack((data))
        data = np.hstack((self.time.reshape(-1, 1), data))
        data = np.vstack((first_row, data))
        df = pd.DataFrame(data)
        df.to_csv(path_to_file[:-3]+'csv', 
            encoding="cp1251", sep=';', float_format='%.3f', decimal=",", header=False, index=False)




    def prepare_figure(self):
        """ Формирует макет осциллограммы.

        """
        max_label_len = max([len(i) for i in self.choosen_signals])

        # макет формируется из строк: 3 строки на аналог, 1 строка на дискрет
        rows = len(self.choosen_signals) * \
            3 + len(self.digitals) + 1

        #
        self.fig = plt.figure(dpi=self.global_settings['IMAGE_DPI'],
                              figsize=(8, rows * 0.3))

        # отступы от края в зависимости от типа файлов
        if self.comtrade:
            self.fig.subplots_adjust(top=0.99,
                                     bottom=0.01,
                                     left=0.21,
                                     right=0.99)
        if self.csv:
            self.fig.subplots_adjust(top=0.99,
                                     bottom=0.01,
                                     left=(0.095 + (max_label_len - 2) * 0.007),
                                     right=0.99)

        self.gs = self.fig.add_gridspec(rows, 1, hspace=0.05)

    def add_analogs(self):
        a = []
        j = 0
        for signal in self.choosen_signals:
            ax = self.fig.add_subplot(self.gs[j * 3:j * 3 + 3, 0])
            j += 1

            def get_num_label_legend(signal):
                """ get_num_label_legend = get_num_label_legend"""
                try:
                    self.analogs[signal]
                except KeyError:
                    return 'err'

                if signal == '':
                    return

                if isinstance(self.analogs[signal], list):
                    num = self.analogs[signal][0]
                    label = signal
                    ylabel = label + f', {self.analogs[signal][1]}'
                else:
                    num = self.analogs[signal]
                    label = signal
                    ylabel = label + f', pu'

                return {'num': num,
                        'label': label,
                        'ylabel': ylabel
                        }

            # получаем данные первого сигнала
            s_1 = get_num_label_legend(signal)

            # если есть дополнительный сигнал в выбранных
            if self.choice[signal] != '':
                # добавляем первый сигнал на участок осциллограммы
                ax.plot(self.time,
                        self.analog_data(s_1['num']),
                        linewidth=0.8,
                        label=s_1['label']
                        )

                # получаем данные второго сигнала
                s_2 = get_num_label_legend(self.choice[signal])

                # если второй сигнал есть в исходном файле, то добавлем его на участок
                if s_2 != 'err':
                    ax.plot(self.time,
                            self.analog_data(s_2['num']),
                            linewidth=0.8,
                            ls='--',
                            label=s_2['label']
                            )
                    ax.legend(loc=0, fontsize='xx-small')

            # если нет дополнительного сигнала
            else:
                ax.plot(self.time,
                        self.analog_data(s_1['num']),
                        linewidth=0.8
                        )

            ax.tick_params(
                axis='y',
                labelsize=6,
                length=0,
                pad=3,
                grid_alpha=0.4)
            ax.tick_params(axis='x', labelsize=6, grid_alpha=0.4)
            plt.ylabel(s_1['ylabel'],
                       fontsize=7,
                       rotation=0,
                       horizontalalignment='right'
                       )
            ax.grid()
            a.append(ax)

    def add_digitals(self):
        a = []
        j = 0
        start_point = len(self.choosen_signals) * 3
        for i in self.digitals:
            ax = self.fig.add_subplot(
                self.gs[start_point + j:start_point + j + 1, 0], ylim=[-0.3, 1.3])
            j += 1
            ax.plot(self.time,
                    self.comtrade.getDigitalChannelData(i),
                    linewidth=0.8)
            ax.set_yticklabels('')
            ax.tick_params(axis='both', labelsize=6, length=0)
            ax.tick_params(axis='y', grid_alpha=0)
            ax.tick_params(axis='x', grid_alpha=0.4)
            plt.ylabel(self.digitals[i], fontsize=6,
                       rotation=0, horizontalalignment='right',
                       verticalalignment='center', wrap=True)
            ax.grid()
            a.append(ax)

    def analogs_select(self, choice):
        """ Выбрать сигналы, из choice, которые есть 
        в self.analogs и сформировать переменную self.choosen_signals
        """
        self.choosen_signals = []
        for signal in choice:
            if signal in self.analogs:
                self.choosen_signals.append(signal)

    def del_empty_digital_signals(self):
        signals = {}
        for i in self.digitals:
            if (1 in self.comtrade.getDigitalChannelData(i) and 0 in self.comtrade.getDigitalChannelData(i)):
                signals.update({i: self.digitals.get(i)})
        self.digitals = signals

    def transform_analogs_names(self):
        """ Убирает номер параметра аналогово сигнала, 
            {1: '80 If', 2: '100 Ug', 3: '120 Fg Hz', ... 24: '829 yTest'}
                ↓↓↓
            {'If': 1, 'Ug': 2, 'Fg': [3, 'Hz'], ... 'yTest': 24}
        """
        new = {}
        for key, val in self.analogs.items():
            val = val.split(' ')
            if len(val) == 3 and self.comtrade:
                new.update({val[1]: [key, val[2]]})
            elif len(val) == 2 and self.comtrade:
                new.update({val[1]: key})

            elif len(val) == 2 and self.csv:
                new.update({val[0]: [key, val[1]]})
            elif len(val) == 1 and self.csv:
                new.update({val[0]: key})
        self.analogs = new


class Helper():
    """docstring for Helper"""

    def __init__(self):
        pass

    def read_settings(self):
        self.settings = self.load_json('settings.txt')
        return self.settings

    def rewrite_settings(self, **add):
        settings = self.read_settings()
        settings.update(add)
        self.write_json(settings, 'settings.txt')

    def load_json(self, name):
        """return {old:new}"""
        with open(name, "r") as file:
            return json.load(file)

    def write_json(self, data, name):
        with open(name, "w") as write_file:
            write_file.write(json.dumps(data, indent=4, ensure_ascii=False))


class Tree():
    """Класс работы с проектом Испытаний"""

    # Используется именно static/... т.к только там можно разместить картинки для превью
    def __init__(self, PROJECT_PATH='static/Испытания'):

        # наличие лишних точек в пути
        self.dots = 0

        self.PROJECT_PATH = PROJECT_PATH

        # чтение настроек, для каких осциллограмм какие параметры выбирать
        self.standard_settings = Helper().load_json('standard_settings.json')

        # получение списка файлов и папок проекта и добавление их в self.project_walk
        self.scan_project()

        # замена лишних точек в имени файлов
        self.replace_dots_in_filenames()

        # замена лишних точек в имени папок
        self.replace_dots_in_foldernames()

        if self.dots == 1:
            # пересканирование проекта
            self.scan_project()
            logging.info(f'Проект пересканирован')

        # получение всех путей к файлам проекта
        self.all_paths = self.get_all_paths()

        # получение тюнсов
        self.tunes = self.get_tunes()

        self.make_tree()

        # print(self.make_tree())
        # self.make_tree()
        # print(list(self.addreses))

        # Helper().write_json(list(enumerate(self.addreses.keys())), 'tree.txt')
        # self.make_images()
        # print(self.addreses.keys())
        # print(self.csv_paths)
        # print(self.cfg_paths)

        # self.write_json(a, "header_replacements.json")
        # print(self.tests_list[140])

    ########### Подготовка ##############

    def scan_project(self):
        """Сканирует папку с проектом
        return: project_walk
        """
        self.project_walk = list(os.walk(self.PROJECT_PATH))

    def replace_dots_in_filenames(self):
        """Переименовывает файлы, где есть лишние точки в 
        имени, отличные от расширения

        """
        for paths in self.project_walk:
            for file in paths[2]:
                if file.count('.') > 1:
                    # корректировка пути
                    path = paths[0].replace('\\', '/')

                    # формирование пути к файлу
                    old_path = f'{path}/{file}'

                    # формирование нового имени файла
                    new_filename = file[:-4].replace('.', '_') + file[-4:]

                    # формирование нового пути
                    new_path = f'{path}/{new_filename}'

                    # переименование
                    os.rename(old_path, new_path)

                    logging.warning(f'Файл {path} переименован в {new_path}')

                    self.dots = 1

    def replace_dots_in_foldernames(self):
        """Переименовывает папки, где есть точки

        """
        for paths in reversed(self.project_walk):
            for folder in paths[1]:
                if folder.count('.') > 0:

                    # корректировка пути
                    path = paths[0].replace('\\', '/')

                    # формирование пути к папке
                    old_path = f'{path}/{folder}'

                    # формирование нового имени
                    new_foldername = folder.replace('.', ',')

                    # формирование нового пути к папке
                    new_path = f'{path}/{new_foldername}'

                    # переименование
                    os.rename(old_path, new_path)

                    logging.warning(
                        f'Папка {old_path} переименована в {new_path}')

                    self.dots = 1

    def get_tunes(self):
        """ Берет файл tunes.txt в файле проекта испытаний и возвращает
        словарь вида:
        {'6': 'swC14@FailPLC', '№параметра':'Имя тюнса', ...}

        """
        def clear(row):
            row = row.replace('\n', '')
            row = row.split('\t')
            row = [i.strip() for i in row]
            row = {row[0]: row[1]}
            return row

        tunes_path = self.PROJECT_PATH + '/tunes.txt'

        if os.path.exists(tunes_path):
            with open(tunes_path, 'r') as f:
                body = {}
                first_line = f.readline().replace('\n', '')
                second_line = f.readline().replace('\n', '')
                head = [first_line, second_line]
                [body.update(clear(i)) for i in f]
            return body

    def get_all_paths(self):
        """ Индексирует все пути
        return: list
        """
        all_paths = []
        self.addreses = {}

        for address, dirs, files in self.project_walk:
            for file in files:
                address = address.replace('\\', '/')

                if file[-3:] == 'cfg' or file[-3:] == 'csv':
                    self.addreses.update({
                        address + '/' + file:
                            {'path': address,
                             'filename': file
                             }})
                    all_paths.append(address + '/' + file)
        return all_paths

    def check_excess_dots(self):
        """ Проверяет наличие двух и более точек в пути файлов.
        Если встречается точка отличная от точки расширения, то комтрейд не
        может обработать файл. Функция создана для предупреждения.
        return: [список файлов с неправильным расширением]
        """
        a = [i for i in self.all_paths if i.count('.') > 1]
        return a

    #####################################

    def make_images(self):
        tree = Helper().load_json('checked.tmp')
        for i in self.addreses:
            choice = tree[i]
            Scope(choice, './' + i)

    def make_tree(self):

        # Начальный номер рисунка
        self.num = Helper().read_settings()['NUM_START']

        logging.info('start make tree')

        tree = self.get_headers()
        logging.info('headers getted')

        checked_for_images = {}

        for header in tree:
            for path in self.addreses:
                if f'/{header}/' in path:
                    all_parameters = self.find_all_parameters(path)
                    cp = self.get_checked_parameters(path)

                    def get_list_part(x, index_):
                        if x is not None:
                            return x[index_]
                        else:
                            return [{'': ''}]

                    def not_checked(x):
                        if x is not None:
                            j = list(x[1][0].keys())
                            difference = all_parameters - set(j)
                            output = {i: "" for i in difference}
                            return output
                            # return all_parameters - set(j)

                        else:
                            return all_parameters

                    def check_img(path):
                        path = path[:-3] + "png"
                        if os.path.exists(path):
                            return path + f"?r={random.random()}"

                    parameters = self.get_all_params_from_path(
                        self.addreses[path]['path'])

                    signature = self.signature(parameters)

                    tree[header]['scopes_paths'].update(
                        {path: {
                            'path': self.addreses[path]['path'],
                            'filename': self.addreses[path]['filename'],
                            'tags': get_list_part(cp, 0),
                            'checked': get_list_part(cp, 1)[0],
                            'not_checked': not_checked(cp),
                            'img': check_img(path),
                            'parameters': parameters,
                            'signature': signature
                        }
                        }
                    )

                    checked_for_images.update({path: get_list_part(cp, 1)[0]})
        Helper().write_json(checked_for_images, "checked.tmp")
        Helper().write_json(tree, "tree.tmp")
        logging.info('Tree is ready!')
        return tree

    def get_checked_parameters(self, path):
        def recursive_find(string, dictionary, j=[]):
            for i in dictionary:
                if i.upper() in string.upper():
                    j.append(i)
                    if isinstance(dictionary[i], dict):
                        return recursive_find(string, dictionary[i], j)
                    else:
                        return j, dictionary[i]
        a = recursive_find(path, self.standard_settings)
        # if a is not None:
        # return a[0]
        return a

    def get_all_params_from_path(self, path):
        upper_path = path.upper()
        source = ''
        replace = ''

        def check_in_path(type_):
            for i in self.standard_settings[type_]:
                if i.upper() in upper_path:
                    return i
            return ''

        # определение типа испытаний
        test = check_in_path('test')
        if test == '':
            if 'ТОЛЧ' in upper_path:
                test = 'Проверка устойчивости'

        # определение регулятора и канала
        regulator = check_in_path('regulator')
        if regulator != '':
            channel = re.search(fr'{regulator}(\d)', path, flags=re.IGNORECASE)
            if channel:
                channel = channel[1]
            # замена на AVR и ECR
            reg_upper = regulator.upper()
            if reg_upper in self.replacements['regulators']:
                regulator = self.replacements['regulators'][reg_upper]
        else:
            channel = ''

        # обработка переходов
        if test == 'переходы':
            jump = re.search(fr'(?:\d+_)*(\w+\d*\b\s*(-\s*\w+\d*\b)+)(\s(.*))*',
                             path, flags=re.IGNORECASE)
            regulator = jump[1]
            replace = jump[4]
            channel = ''

        # обработка Неуспешного НВ
        if 'НВ' in test:
            jump = re.search(fr'(?:\d+_)*(\w+)(\d)\b\s*([~=])*\s*',
                             path, flags=re.IGNORECASE)
            regulator = jump[1]
            channel = jump[2]
            if jump[3] == "~":
                source = "от источника переменного тока"
            if jump[3] == "=":
                source = "от источника постоянного тока"


        # определение амплитуды толчка
        pulse = re.search(r'[+,-]\d+%(?:\s*\D+\b)*[$]*', path)
        if pulse:
            pulse = pulse[0].strip()
        else:
            pulse = ''

        # определение длительности толчка
        pulse_time = re.search(r'\d+\s*(сек|sec)', path)
        if pulse_time:
            pulse_time = pulse_time[0]
        else:
            pulse_time = ''

        # определение параметров, записанных в пути
        tunes = re.findall(r'[TТV]\d+\s*\w+?\s*=\s*[\d,]+', path)

        def add_tunes_name(parameter):
            p_list = parameter.split('=')

            if self.tunes is not None and p_list[0][1:] in self.tunes:
                p_list.insert(1, self.tunes[p_list[0][1:]])
                new_param = p_list[0] + p_list[1] + ' = ' + p_list[2]
                parameter = new_param

            # удаление запятой после тюнса
            if parameter.endswith(','):
                parameter = parameter[:-1]

            return parameter

        if tunes is not None:
            tunes = [add_tunes_name(i) for i in tunes]
        if '_set'.upper() in upper_path:
            set_ = 'set'
        else:
            set_ = ''

        # формирование выходного словаря
        output = {
            'test': test,
            'regulator': regulator,
            'channel': channel,
            'pulse': pulse,
            'pulse_time': pulse_time,
            'tunes': tunes,
            'source': source,
            'replace': replace,
            'set_': set_
        }

        return output

    def find_all_parameters(self, path=False):
        """ Ищет все уникальные аналоговые сигналы в папке с проектом
        return: {a,b,c}
        """
        parameters = set()
        if path is False:
            paths = self.get_all_paths()
        else:
            paths = [path, ]
        for i in paths:
            path = './' + i
            if i[-3:] == 'cfg':
                is_csv = False
                comtradeObj = ComtradeRecord(path)
                comtradeObj.ReadCFG()
                parameters = parameters.union(
                    set(list(comtradeObj.signals['Analog'].values())), parameters)

            if i[-3:] == 'csv':
                is_csv = True
                csv_obj = Csv(path)
                parameters = parameters.union(
                    set(list(csv_obj.signals.values())))

        parameters_clear = set()
        for i in parameters:
            a = i.split(' ')

            if is_csv is True and len(a) > 1:
                parameters_clear = parameters_clear.union(
                    (a[0],), parameters_clear)
            elif len(a) > 1:
                parameters_clear = parameters_clear.union(
                    (a[1],), parameters_clear)
            else:
                parameters_clear = parameters_clear.union(
                    (a[0],), parameters_clear)
        return parameters_clear

    def get_headers(self):
        """ Подготовка заголовков на основе второго уровня.
        ./Испытания/Заголовок
        return: list[self.headers]
        """
        raw_headers = self.project_walk[0][1]

        def clear(a):
            a = a.split('_')
            return a[1]

        headers = [clear(i) for i in raw_headers]
        headers = dict(zip(raw_headers, self.repalce_using_dict(headers)))

        for i in headers:
            headers[i] = {'change_name': headers[i], 'scopes_paths': {}}
        return headers

    def repalce_using_dict(self, data_list):
        """Замена частей заголовков используя словарь header_replacements.json
        return: list
        """
        try:
            self.replacements = Helper().load_json("header_replacements.json")
        except json.decoder.JSONDecodeError:
            print(
                "Не выполнено. Проверьте файл замен, \
                возможно есть запятая в последней строке")
            exit()

        def change_str(old):
            for key, new in self.replacements.items():
                if key in old:
                    return old.replace(key, new)
            return old

        return [change_str(i) for i in data_list]

    def signature(self, parameters):
        source = parameters['source']

        replacements = self.standard_settings['replacements']

        try:
            regulator = replacements[parameters['regulator']]
        except:
            regulator = parameters['regulator']

        try:
            replace = replacements[parameters['replace']]
            replace = ' ' + replace
        except:
            replace = ''

        test = f"f\"{replacements[parameters['test']]}\""
        channel = parameters['channel']
        pulse = parameters['pulse']

        if parameters['regulator'] == 'AVR':
            key_parameter = 'Ug'
        elif parameters['regulator'] == 'ECR':
            key_parameter = 'Ie'
        elif parameters['regulator'] == 'CosPhi':
            key_parameter = 'CosPhi'
        elif parameters['regulator'] == 'Qg':
            key_parameter = 'Qg'

        if parameters['pulse_time']:
            pulse_time = f" длительностью {parameters['pulse_time']}"
        else:
            pulse_time = ""

        if parameters['tunes']:
            tunes = " "+", ".join(parameters['tunes'])
        else:
            tunes = ""

        if parameters['set_']:
            set_ = f"\n\nВ результате проверки установлены следующие параметры: {tunes}"
        else:
            set_ = ""

        body = f"Рисунок {self.num} - {eval(test)}{set_}"

        self.num += 1

        return body


def make_protocol():
    tree = Helper().load_json('tree.tmp')
    document = Document('template.docx')

    for key, branch in tree.items():
        document.add_heading(branch['change_name'], 0)
        for full_path, params in branch['scopes_paths'].items():

            if params['img'] is not None:
                img_path = params['img'].split("?r=")[0]
                document.add_picture(f"./{img_path}", width=Cm(16))

            if params['signature'] != "":
                document.add_paragraph(params['signature'])
                document.add_paragraph('')

    document.save('./static/protocol.docx')


def main():

    pass



if __name__ == '__main__':
    main()

